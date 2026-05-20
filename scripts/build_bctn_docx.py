from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "BCTN_GearZone_PhanTichThietKe.md"
DOCX_PATH = ROOT / "BCTN_GearZone_PhanTichThietKe.docx"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title_page(doc: Document) -> None:
    for text, size, bold in [
        ("TRƯỜNG ĐẠI HỌC KĨ THUẬT CÔNG NGHIỆP", 14, True),
        ("KHOA ĐIỆN TỬ", 13, True),
        ("---oOo---", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ĐỒ ÁN TỐT NGHIỆP")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(20)

    for text in [
        "NGÀNH: KĨ THUẬT MÁY TÍNH",
        "HỆ: ĐẠI HỌC CHÍNH QUY",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ĐỀ TÀI: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG WEBSITE BÁN PHỤ KIỆN MÁY TÍNH GAMING TRỰC TUYẾN - GEARZONE")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)

    doc.add_paragraph()
    meta = [
        ("GIÁO VIÊN HƯỚNG DẪN", "....................................."),
        ("SINH VIÊN THỰC HIỆN", "....................................."),
        ("LỚP", "....................................."),
        ("MSSV", "....................................."),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(3.5)
        r = p.add_run(f"{label}: {value}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)

    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THÁI NGUYÊN - NĂM 2026")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    doc.add_page_break()


def add_assignment_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRƯỜNG ĐHKTCN        CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KHOA ĐIỆN TỬ                 Độc lập - Tự do - Hạnh phúc")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NHIỆM VỤ THIẾT KẾ TỐT NGHIỆP")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)
    items = [
        "Sinh viên: .....................................        MSSV: .....................................",
        "Lớp: .....................................        Khoá: .....................................",
        "Ngành học: Kỹ thuật máy tính",
        "Giáo viên hướng dẫn: .....................................",
        "1. Tên đề tài tốt nghiệp: Phân tích và thiết kế hệ thống website bán phụ kiện máy tính gaming trực tuyến - GearZone.",
        "2. Các số liệu ban đầu: mã nguồn website GearZone, cơ sở dữ liệu Prisma/SQLite, tài liệu bài tập phân tích thiết kế hướng đối tượng.",
        "3. Nội dung các phần thuyết minh và tính toán:",
        "- Khảo sát hiện trạng và xác định yêu cầu.",
        "- Phân tích các chức năng nghiệp vụ và tác nhân của hệ thống.",
        "- Thiết kế sơ đồ UML, cơ sở dữ liệu và kiến trúc triển khai.",
        "- Cài đặt website, kiểm thử các chức năng chính và đánh giá kết quả.",
        "4. Số lượng phần mềm, bảng biểu, bản vẽ: báo cáo thuyết minh, hệ thống website, các sơ đồ UML và bảng kiểm thử.",
        "5. Ngày giao nhiệm vụ: ...... / ...... / 2026",
        "6. Ngày hoàn thành nhiệm vụ: ...... / ...... / 2026",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.style = doc.styles["Normal"]
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, txt in zip(t.rows[0].cells, ["BCN KHOA", "TRƯỞNG BỘ MÔN", "GIÁO VIÊN HƯỚNG DẪN"]):
        set_cell_text(cell, txt, True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_score_page(doc: Document) -> None:
    lines = [
        "TRƯỜNG ĐHKTCN        CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "KHOA ĐIỆN TỬ                 Độc lập - Tự do - Hạnh phúc",
        "",
        "PHIẾU GHI ĐIỂM",
        "HƯỚNG DẪN ĐỒ ÁN TỐT NGHIỆP",
        "",
        "Sinh viên: .....................................",
        "Lớp: .....................................",
        "GVHD: .....................................",
        "Đề tài: Phân tích và thiết kế hệ thống website bán phụ kiện máy tính gaming trực tuyến - GearZone.",
        "",
        "NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN",
    ]
    for line in lines:
        p = doc.add_paragraph()
        if line in {"PHIẾU GHI ĐIỂM", "HƯỚNG DẪN ĐỒ ÁN TỐT NGHIỆP", "NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN"}:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = line.isupper() and bool(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(13 if line in {"PHIẾU GHI ĐIỂM", "HƯỚNG DẪN ĐỒ ÁN TỐT NGHIỆP"} else 12)
    for _ in range(7):
        doc.add_paragraph("................................................................................................................")
    doc.add_paragraph("Xếp loại: ....................        Điểm: ....................")
    p = doc.add_paragraph("Thái Nguyên, ngày .... tháng .... năm 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph("GIÁO VIÊN HƯỚNG DẪN")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()


def markdown_content() -> str:
    return r'''# TRƯỜNG ĐẠI HỌC KĨ THUẬT CÔNG NGHIỆP

## KHOA ĐIỆN TỬ

---oOo---

# ĐỒ ÁN TỐT NGHIỆP

**NGÀNH:** KĨ THUẬT MÁY TÍNH

**HỆ:** ĐẠI HỌC CHÍNH QUY

**ĐỀ TÀI:** PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG WEBSITE BÁN PHỤ KIỆN MÁY TÍNH GAMING TRỰC TUYẾN - GEARZONE

**GIÁO VIÊN HƯỚNG DẪN:** .....................................

**SINH VIÊN THỰC HIỆN:** .....................................

**LỚP:** .....................................

**MSSV:** .....................................

**THÁI NGUYÊN - NĂM 2026**

---

# NHIỆM VỤ THIẾT KẾ TỐT NGHIỆP

**Sinh viên:** .....................................  
**MSSV:** .....................................  
**Lớp:** .....................................  
**Khoá:** .....................................  
**Ngành học:** Kỹ thuật máy tính  
**Giáo viên hướng dẫn:** .....................................

1. Tên đề tài tốt nghiệp: Phân tích và thiết kế hệ thống website bán phụ kiện máy tính gaming trực tuyến - GearZone.
2. Các số liệu ban đầu: mã nguồn website GearZone, cơ sở dữ liệu Prisma/SQLite, tài liệu bài tập phân tích thiết kế hướng đối tượng.
3. Nội dung các phần thuyết minh và tính toán:

- Khảo sát hiện trạng và xác định yêu cầu.
- Phân tích các chức năng nghiệp vụ và tác nhân của hệ thống.
- Thiết kế sơ đồ UML, cơ sở dữ liệu và kiến trúc triển khai.
- Cài đặt website, kiểm thử các chức năng chính và đánh giá kết quả.

4. Số lượng phần mềm, bảng biểu, bản vẽ: báo cáo thuyết minh, hệ thống website, các sơ đồ UML và bảng kiểm thử.
5. Ngày giao nhiệm vụ: ...... / ...... / 2026.
6. Ngày hoàn thành nhiệm vụ: ...... / ...... / 2026.

---

# PHIẾU GHI ĐIỂM

## HƯỚNG DẪN ĐỒ ÁN TỐT NGHIỆP

**Sinh viên:** .....................................  
**Lớp:** .....................................  
**GVHD:** .....................................  
**Đề tài:** Phân tích và thiết kế hệ thống website bán phụ kiện máy tính gaming trực tuyến - GearZone.

## NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN

................................................................................................................

................................................................................................................

................................................................................................................

................................................................................................................

**Xếp loại:** ....................        **Điểm:** ....................

---

# BÁO CÁO THỰC TẬP TỐT NGHIỆP

## Đề tài: Phân tích và thiết kế hệ thống website bán phụ kiện máy tính gaming trực tuyến - GearZone

**Giảng viên hướng dẫn:** .....................................

**Sinh viên thực hiện:** .....................................

**Mã sinh viên:** .....................................

**Lớp:** .....................................

**Địa điểm:** Thái Nguyên, năm 2026

---

# LỜI CẢM ƠN

Em xin gửi lời cảm ơn chân thành tới thầy/cô ..................................... đã tận tình hướng dẫn, góp ý và định hướng cho em trong quá trình thực hiện đề tài. Những chỉ dẫn của thầy/cô giúp em hiểu rõ hơn cách khảo sát yêu cầu, phân tích hệ thống theo hướng đối tượng và trình bày kết quả nghiên cứu thành một báo cáo hoàn chỉnh.

Em cũng xin cảm ơn quý thầy cô khoa Điện tử và bộ môn Công nghệ thông tin đã truyền đạt các kiến thức nền tảng về lập trình web, cơ sở dữ liệu, công nghệ phần mềm và phân tích thiết kế hệ thống. Đây là cơ sở quan trọng để em vận dụng vào việc xây dựng website GearZone.

Trong quá trình thực hiện, do thời gian và kinh nghiệm còn hạn chế nên báo cáo khó tránh khỏi thiếu sót. Em rất mong nhận được ý kiến đóng góp của quý thầy cô để hoàn thiện sản phẩm và rút kinh nghiệm cho các dự án sau.

Em xin chân thành cảm ơn!

# LỜI NÓI ĐẦU

Trong những năm gần đây, thương mại điện tử phát triển mạnh mẽ và dần trở thành kênh mua sắm quen thuộc của người tiêu dùng. Cùng với đó, cộng đồng game thủ và người dùng máy tính hiệu năng cao ngày càng quan tâm đến các thiết bị như chuột gaming, bàn phím cơ, tai nghe, màn hình tần số quét cao và phụ kiện setup góc làm việc. Nhu cầu mua sắm trực tuyến các sản phẩm này đòi hỏi một hệ thống có thông tin rõ ràng, hình ảnh trực quan, khả năng tìm kiếm nhanh, giỏ hàng thuận tiện và quy trình đặt hàng minh bạch.

Đề tài "Phân tích và thiết kế hệ thống website bán phụ kiện máy tính gaming trực tuyến - GearZone" được thực hiện nhằm xây dựng một website thương mại điện tử chuyên về gaming gear. Hệ thống hỗ trợ khách hàng xem sản phẩm, tìm kiếm, thêm vào giỏ hàng, đặt hàng, theo dõi đơn hàng; đồng thời hỗ trợ quản trị viên quản lý sản phẩm, đơn hàng, người dùng và cấu hình website.

Báo cáo trình bày quá trình khảo sát, phân tích yêu cầu, thiết kế hệ thống bằng UML, thiết kế cơ sở dữ liệu, thực hiện website và thử nghiệm các chức năng chính. Cấu trúc báo cáo được trình bày theo mẫu đồ án gồm sáu chương: giới thiệu, phương pháp thực hiện, thiết kế, thực hiện website, thử nghiệm và kết luận.

# MỤC LỤC

- NHIỆM VỤ THIẾT KẾ TỐT NGHIỆP
- PHIẾU GHI ĐIỂM
- LỜI CẢM ƠN
- LỜI NÓI ĐẦU
- DANH MỤC HÌNH VẼ
- CHƯƠNG I: GIỚI THIỆU
- CHƯƠNG II: PHƯƠNG PHÁP THỰC HIỆN
- CHƯƠNG III: THIẾT KẾ
- CHƯƠNG IV: THỰC HIỆN WEBSITE
- CHƯƠNG V: THỬ NGHIỆM
- CHƯƠNG VI: KẾT LUẬN
- TÀI LIỆU THAM KHẢO

# DANH MỤC HÌNH VẼ

| STT | Hình ảnh | Nội dung |
|---|---|---|
| 1 | Hình 2-1 | Sơ đồ chức năng hệ thống GearZone |
| 2 | Hình 2-2 | Sơ đồ use case tổng quát |
| 3 | Hình 3-1 | Sơ đồ lớp của hệ thống |
| 4 | Hình 3-2 | Sơ đồ tuần tự đặt hàng |
| 5 | Hình 3-3 | Sơ đồ trạng thái đơn hàng |
| 6 | Hình 3-4 | Sơ đồ hoạt động mua hàng |
| 7 | Hình 3-5 | Sơ đồ thành phần hệ thống |
| 8 | Hình 3-6 | Sơ đồ triển khai hệ thống |

# CHƯƠNG I: GIỚI THIỆU

## 1.1 Mục tiêu của đề tài

Mục tiêu của đề tài là phân tích, thiết kế và xây dựng website GearZone phục vụ hoạt động bán phụ kiện máy tính gaming trực tuyến. Website cần đáp ứng nhu cầu mua sắm của khách hàng và nhu cầu quản trị vận hành của cửa hàng.

Các mục tiêu cụ thể gồm:

- Khảo sát các website bán thiết bị công nghệ, gaming gear và rút ra yêu cầu cho hệ thống.
- Xác định chức năng chính của khách hàng, quản trị viên và hệ thống thanh toán.
- Thiết kế các mô hình UML gồm use case, class, sequence, state, activity, component và deployment.
- Thiết kế cơ sở dữ liệu lưu trữ người dùng, sản phẩm, danh mục, giỏ hàng, đơn hàng và cấu hình website.
- Xây dựng website bằng Next.js, React, TypeScript, Prisma ORM, SQLite và TailwindCSS.
- Kiểm thử các chức năng chính, đánh giá kết quả và đề xuất hướng phát triển.

## 1.2 Những thách thức cần giải quyết

Website thương mại điện tử cho nhóm sản phẩm gaming gear có một số thách thức riêng. Người mua thường quan tâm đến thông số kỹ thuật chi tiết như switch bàn phím, DPI chuột, tần số quét màn hình, chuẩn kết nối, độ trễ, bảo hành và khả năng tương thích. Vì vậy hệ thống phải trình bày thông tin sản phẩm rõ ràng, có ảnh minh họa và khả năng lọc/tìm kiếm phù hợp.

Ngoài ra, quy trình mua hàng cần hạn chế lỗi tồn kho, đảm bảo giỏ hàng cập nhật chính xác và đơn hàng có trạng thái rõ ràng. Phần quản trị phải đủ đơn giản để thêm sản phẩm, cập nhật hình ảnh, chỉnh tồn kho, theo dõi đơn hàng và quản lý người dùng.

## 1.3 Nội dung, phạm vi thực hiện

Phạm vi đề tài tập trung vào website thương mại điện tử B2C cho cửa hàng GearZone. Đối tượng sử dụng gồm khách hàng và quản trị viên.

Nội dung thực hiện gồm:

- Giao diện người dùng: trang chủ, danh sách sản phẩm, chi tiết sản phẩm, giỏ hàng, đăng nhập, đăng ký và đơn hàng.
- Quản trị hệ thống: dashboard, quản lý sản phẩm, quản lý đơn hàng, quản lý người dùng, cấu hình banner/video và thông tin liên hệ.
- Cơ sở dữ liệu: thiết kế các bảng users, products, categories, carts, cart_items, orders, order_items, reviews và settings.
- Bảo mật: mã hóa mật khẩu bằng bcryptjs, xác thực JWT, phân quyền USER/ADMIN.
- Kiểm thử: kiểm thử chức năng đăng ký, đăng nhập, tìm kiếm sản phẩm, thêm giỏ hàng, đặt hàng và quản trị.

## 1.4 Kết quả cần đạt được

| Chức năng | Kết quả cần đạt |
|---|---|
| Xem sản phẩm | Hiển thị danh sách sản phẩm, ảnh, giá, danh mục, trạng thái còn hàng |
| Tìm kiếm/lọc | Tìm theo tên, mô tả, danh mục; hỗ trợ sắp xếp theo giá hoặc sản phẩm mới |
| Chi tiết sản phẩm | Hiển thị ảnh, mô tả, thông số kỹ thuật, đánh giá và nút thêm giỏ hàng |
| Giỏ hàng | Thêm, xóa, cập nhật số lượng, tính tổng tiền |
| Đặt hàng | Tạo đơn hàng, lưu thông tin nhận hàng, chọn phương thức thanh toán |
| Quản lý sản phẩm | Thêm, sửa, xóa, upload ảnh và cập nhật tồn kho |
| Quản lý đơn hàng | Xem danh sách, chi tiết đơn và cập nhật trạng thái |
| Quản lý người dùng | Xem danh sách người dùng và phân quyền |

# CHƯƠNG II: PHƯƠNG PHÁP THỰC HIỆN

## 2.1 Các hệ thống tương tự

Trước khi xây dựng GearZone, đề tài khảo sát một số hệ thống bán thiết bị công nghệ và phụ kiện gaming phổ biến để nhận diện các chức năng cần có.

### 2.1.1 GearVN

GearVN là website chuyên bán máy tính, laptop và thiết bị gaming. Ưu điểm là danh mục rõ ràng, sản phẩm có nhiều ảnh, thông số chi tiết, có chính sách bảo hành và khuyến mãi nổi bật. Hạn chế là giao diện có nhiều banner, mật độ thông tin cao nên người dùng mới có thể mất thời gian tìm đúng sản phẩm.

### 2.1.2 Phong Vũ

Phong Vũ có hệ sinh thái sản phẩm lớn, hỗ trợ nhiều nhóm hàng công nghệ và có uy tín thương hiệu. Điểm mạnh là hệ thống danh mục rộng, tìm kiếm tương đối tốt và có cửa hàng vật lý. Điểm hạn chế với phân khúc gaming gear là chưa tạo cảm giác chuyên biệt cho cộng đồng game thủ.

### 2.1.3 HACOM

HACOM có ưu thế về giá bán, sản phẩm đa dạng và nhiều chương trình khuyến mãi. Tuy nhiên giao diện nhiều thông tin, bố cục còn dày và chưa tối ưu cho trải nghiệm mua nhanh trên thiết bị di động.

Từ khảo sát trên, GearZone định hướng giao diện tập trung, hiện đại, giảm nhiễu thông tin, ưu tiên tìm kiếm nhanh, thông tin sản phẩm rõ và quy trình đặt hàng ngắn.

## 2.2 Công nghệ sử dụng

| Công nghệ | Vai trò |
|---|---|
| Next.js 15 | Framework xây dựng giao diện và API route |
| React 19 | Xây dựng component giao diện |
| TypeScript | Tăng tính an toàn khi phát triển |
| Prisma ORM | Làm việc với cơ sở dữ liệu theo mô hình object-relational mapping |
| SQLite | Cơ sở dữ liệu thử nghiệm trong phạm vi đề tài |
| TailwindCSS | Thiết kế giao diện nhanh, responsive |
| bcryptjs | Mã hóa mật khẩu |
| jose/JWT | Xác thực phiên đăng nhập |
| SWR | Lấy dữ liệu phía client |

## 2.3 Phân tích yêu cầu

### 2.3.1 Các quy trình nghiệp vụ

Quy trình khách hàng mua sản phẩm gồm các bước: truy cập website, tìm kiếm hoặc duyệt danh mục, xem chi tiết sản phẩm, thêm sản phẩm vào giỏ hàng, kiểm tra giỏ hàng, nhập thông tin giao hàng, chọn phương thức thanh toán và xác nhận đặt hàng.

Quy trình quản trị gồm: đăng nhập admin, thêm/cập nhật sản phẩm, kiểm tra đơn hàng mới, xác nhận đơn, cập nhật trạng thái xử lý/giao hàng, quản lý người dùng và cấu hình nội dung website.

### 2.3.2 Sơ đồ chức năng

```mermaid
flowchart TD
    A[Hệ thống GearZone] --> B[Quản lý người dùng]
    A --> C[Quản lý sản phẩm]
    A --> D[Quản lý giỏ hàng]
    A --> E[Quản lý đơn hàng]
    A --> F[Quản trị hệ thống]
    B --> B1[Đăng ký]
    B --> B2[Đăng nhập/Đăng xuất]
    B --> B3[Cập nhật thông tin cá nhân]
    C --> C1[Xem danh sách]
    C --> C2[Tìm kiếm và lọc]
    C --> C3[Xem chi tiết]
    C --> C4[Đánh giá sản phẩm]
    D --> D1[Thêm vào giỏ]
    D --> D2[Cập nhật số lượng]
    D --> D3[Xóa sản phẩm]
    E --> E1[Tạo đơn hàng]
    E --> E2[Thanh toán]
    E --> E3[Theo dõi trạng thái]
    F --> F1[CRUD sản phẩm]
    F --> F2[Quản lý đơn hàng]
    F --> F3[Quản lý người dùng]
    F --> F4[Cấu hình website]
```

Hình 2-1: Sơ đồ chức năng hệ thống GearZone.

### 2.3.3 Sơ đồ use case tổng quát

```mermaid
flowchart LR
    KH[Khách hàng] --> UC1((Đăng ký/Đăng nhập))
    KH --> UC2((Tìm kiếm sản phẩm))
    KH --> UC3((Xem chi tiết sản phẩm))
    KH --> UC4((Quản lý giỏ hàng))
    KH --> UC5((Đặt hàng))
    KH --> UC6((Theo dõi đơn hàng))
    KH --> UC7((Đánh giá sản phẩm))
    AD[Quản trị viên] --> UC8((Quản lý sản phẩm))
    AD --> UC9((Quản lý đơn hàng))
    AD --> UC10((Quản lý người dùng))
    AD --> UC11((Cấu hình hệ thống))
    PAY[Cổng thanh toán] --> UC12((Xác nhận thanh toán))
    UC5 --> UC12
```

Hình 2-2: Sơ đồ use case tổng quát.

# CHƯƠNG III: THIẾT KẾ

## 3.1 Kiến trúc hệ thống

GearZone được thiết kế theo kiến trúc web ba lớp. Lớp giao diện sử dụng Next.js/React để hiển thị trang khách hàng và trang quản trị. Lớp xử lý nghiệp vụ nằm trong API route và các module service, chịu trách nhiệm xác thực, xử lý giỏ hàng, đặt hàng và quản trị. Lớp dữ liệu sử dụng Prisma ORM kết nối SQLite trong môi trường thử nghiệm.

## 3.2 Thiết kế cơ sở dữ liệu

Các thực thể chính gồm User, Product, Category, Cart, CartItem, Order, OrderItem, Review và Setting.

| Bảng | Ý nghĩa | Thuộc tính chính |
|---|---|---|
| users | Lưu tài khoản người dùng | id, name, email, passwordHash, role, createdAt |
| categories | Lưu danh mục sản phẩm | id, name, slug, description |
| products | Lưu sản phẩm | id, name, price, imageUrl, stock, description, categoryId |
| carts | Lưu giỏ hàng | id, userId, createdAt |
| cart_items | Chi tiết giỏ hàng | id, cartId, productId, quantity |
| orders | Đơn hàng | id, userId, total, status, paymentMethod, address |
| order_items | Chi tiết đơn hàng | id, orderId, productId, quantity, price |
| reviews | Đánh giá sản phẩm | id, userId, productId, rating, comment |
| settings | Cấu hình website | id, key, value |

## 3.3 Sơ đồ lớp

```mermaid
classDiagram
    class User {
        +String id
        +String name
        +String email
        +String passwordHash
        +String role
        +login()
        +logout()
    }
    class Product {
        +String id
        +String name
        +Float price
        +Int stock
        +String description
        +updateStock()
    }
    class Category {
        +String id
        +String name
        +String slug
    }
    class Cart {
        +String id
        +addItem()
        +removeItem()
        +calculateTotal()
    }
    class CartItem {
        +Int quantity
        +Float unitPrice
    }
    class Order {
        +String id
        +Float total
        +String status
        +create()
        +cancel()
        +updateStatus()
    }
    class OrderItem {
        +Int quantity
        +Float price
    }
    class Review {
        +Int rating
        +String comment
    }
    User "1" --> "0..1" Cart
    User "1" --> "0..*" Order
    User "1" --> "0..*" Review
    Category "1" --> "0..*" Product
    Cart "1" --> "1..*" CartItem
    Product "1" --> "0..*" CartItem
    Order "1" --> "1..*" OrderItem
    Product "1" --> "0..*" OrderItem
    Product "1" --> "0..*" Review
```

Hình 3-1: Sơ đồ lớp của hệ thống GearZone.

## 3.4 Sơ đồ tuần tự đặt hàng

```mermaid
sequenceDiagram
    actor KH as Khách hàng
    participant UI as Giao diện giỏ hàng
    participant API as Order API
    participant DB as Cơ sở dữ liệu
    participant PAY as Cổng thanh toán
    KH->>UI: Kiểm tra giỏ hàng
    UI->>API: Gửi thông tin đặt hàng
    API->>DB: Kiểm tra tồn kho và tạo đơn
    DB-->>API: Trả mã đơn hàng
    API->>PAY: Khởi tạo thanh toán nếu cần
    PAY-->>API: Trả kết quả thanh toán
    API->>DB: Cập nhật trạng thái đơn
    API-->>UI: Trả thông báo thành công
    UI-->>KH: Hiển thị xác nhận đơn hàng
```

Hình 3-2: Sơ đồ tuần tự đặt hàng.

## 3.5 Sơ đồ trạng thái đơn hàng

```mermaid
stateDiagram-v2
    [*] --> ChoXacNhan
    ChoXacNhan --> DangXuLy: Admin xác nhận
    ChoXacNhan --> DaHuy: Khách hủy/Admin hủy
    DangXuLy --> DangGiao: Bàn giao vận chuyển
    DangGiao --> DaGiao: Giao thành công
    DaGiao --> HoanThanh: Khách xác nhận
    DangXuLy --> DaHuy: Hủy trước khi giao
    DaHuy --> [*]
    HoanThanh --> [*]
```

Hình 3-3: Sơ đồ trạng thái đơn hàng.

## 3.6 Sơ đồ hoạt động mua hàng

```mermaid
flowchart TD
    A([Bắt đầu]) --> B[Truy cập GearZone]
    B --> C[Tìm kiếm hoặc chọn danh mục]
    C --> D[Xem chi tiết sản phẩm]
    D --> E{Muốn mua?}
    E -- Không --> C
    E -- Có --> F[Thêm vào giỏ hàng]
    F --> G{Tiếp tục mua?}
    G -- Có --> C
    G -- Không --> H[Kiểm tra giỏ hàng]
    H --> I[Nhập thông tin giao hàng]
    I --> J[Chọn phương thức thanh toán]
    J --> K[Xác nhận đặt hàng]
    K --> L[Theo dõi trạng thái đơn]
    L --> M([Kết thúc])
```

Hình 3-4: Sơ đồ hoạt động mua hàng.

## 3.7 Sơ đồ thành phần

```mermaid
flowchart TB
    Browser[Web Browser] --> Next[Next.js App]
    Next --> Storefront[Storefront Components]
    Next --> Admin[Admin Components]
    Next --> API[API Routes]
    API --> Auth[Auth Module]
    API --> Product[Product Module]
    API --> Order[Order Module]
    API --> Setting[Setting Module]
    Auth --> Prisma[Prisma ORM]
    Product --> Prisma
    Order --> Prisma
    Setting --> Prisma
    Prisma --> DB[(SQLite Database)]
    API --> Upload[Image Upload]
```

Hình 3-5: Sơ đồ thành phần hệ thống.

## 3.8 Sơ đồ triển khai

```mermaid
flowchart TB
    User[Người dùng] --> Browser[Trình duyệt]
    Browser --> Web[Next.js Server]
    Web --> API[API Routes]
    API --> DB[(SQLite/Database)]
    API --> Storage[Thư mục public/uploads]
    Admin[Quản trị viên] --> Browser
```

Hình 3-6: Sơ đồ triển khai hệ thống.

# CHƯƠNG IV: THỰC HIỆN WEBSITE

## 4.1 Môi trường phát triển

Website được thực hiện trong môi trường Node.js với framework Next.js. Mã nguồn sử dụng TypeScript, TailwindCSS và Prisma. Cơ sở dữ liệu thử nghiệm là SQLite để thuận tiện triển khai trong phạm vi đồ án.

Các thư mục chính của dự án:

- `src/app`: chứa các page, layout và API route của Next.js.
- `src/components`: chứa các component giao diện dùng lại.
- `src/lib`: chứa cấu hình database, auth và tiện ích.
- `prisma`: chứa schema cơ sở dữ liệu và seed dữ liệu.
- `public/uploads`: lưu ảnh sản phẩm và tài nguyên upload.

## 4.2 Thực hiện giao diện người dùng

Trang người dùng gồm trang chủ, trang sản phẩm, trang chi tiết sản phẩm, giỏ hàng, đăng nhập, đăng ký và đơn hàng. Giao diện được thiết kế responsive để dùng được trên máy tính và thiết bị di động. Các thành phần như thanh điều hướng, catalog sản phẩm, gallery ảnh, tab thông số và nút thêm giỏ hàng được tách thành component riêng.

Luồng mua hàng được tối ưu theo hướng ngắn gọn: người dùng xem sản phẩm, chọn số lượng, thêm vào giỏ, kiểm tra giỏ hàng và tạo đơn. Khi chưa đăng nhập, hệ thống yêu cầu đăng nhập để đảm bảo đơn hàng gắn với tài khoản.

## 4.3 Thực hiện trang quản trị

Trang quản trị dành cho tài khoản ADMIN. Các chức năng chính gồm:

- Dashboard tổng quan số lượng sản phẩm, người dùng và đơn hàng.
- Quản lý sản phẩm: thêm, sửa, xóa, cập nhật tồn kho và ảnh.
- Quản lý đơn hàng: xem chi tiết và cập nhật trạng thái.
- Quản lý người dùng: xem danh sách và vai trò.
- Cấu hình hệ thống: cập nhật thông tin liên hệ, banner/video.

## 4.4 Xác thực và phân quyền

Hệ thống sử dụng bcryptjs để mã hóa mật khẩu trước khi lưu. Khi người dùng đăng nhập thành công, API tạo JWT để lưu phiên xác thực. Middleware kiểm tra token và quyền truy cập trước khi cho phép vào các trang cần đăng nhập hoặc trang quản trị.

Phân quyền gồm hai vai trò chính: USER và ADMIN. USER được sử dụng các chức năng mua hàng, còn ADMIN được truy cập khu vực quản trị.

## 4.5 Kết nối cơ sở dữ liệu

Prisma ORM giúp định nghĩa schema dữ liệu, sinh client truy vấn và giảm lỗi khi thao tác với database. Các API route sử dụng Prisma để đọc/ghi dữ liệu sản phẩm, người dùng, giỏ hàng và đơn hàng.

# CHƯƠNG V: THỬ NGHIỆM

## 5.1 Mục tiêu thử nghiệm

Mục tiêu thử nghiệm là kiểm tra các chức năng chính có hoạt động đúng yêu cầu hay không, dữ liệu có được lưu chính xác hay không và các tình huống lỗi cơ bản có được xử lý phù hợp hay không.

## 5.2 Bảng kiểm thử chức năng

| STT | Chức năng | Dữ liệu kiểm thử | Kết quả mong đợi | Kết quả |
|---|---|---|---|---|
| 1 | Đăng ký | Email hợp lệ, mật khẩu hợp lệ | Tạo tài khoản mới | Đạt |
| 2 | Đăng nhập | Email/mật khẩu đúng | Trả phiên đăng nhập | Đạt |
| 3 | Đăng nhập sai | Mật khẩu sai | Hiển thị lỗi xác thực | Đạt |
| 4 | Xem sản phẩm | Truy cập trang sản phẩm | Hiển thị danh sách sản phẩm | Đạt |
| 5 | Tìm kiếm | Nhập từ khóa sản phẩm | Trả sản phẩm phù hợp | Đạt |
| 6 | Thêm giỏ hàng | Chọn sản phẩm còn hàng | Sản phẩm xuất hiện trong giỏ | Đạt |
| 7 | Cập nhật giỏ hàng | Tăng/giảm số lượng | Tổng tiền cập nhật | Đạt |
| 8 | Đặt hàng | Thông tin giao hàng hợp lệ | Tạo đơn hàng mới | Đạt |
| 9 | Quản lý sản phẩm | Admin thêm sản phẩm | Sản phẩm được lưu | Đạt |
| 10 | Cập nhật đơn hàng | Admin đổi trạng thái | Trạng thái mới được lưu | Đạt |

## 5.3 Đánh giá kết quả thử nghiệm

Qua thử nghiệm, các chức năng chính của hệ thống hoạt động đúng với yêu cầu ban đầu. Giao diện có thể truy cập trên các kích thước màn hình phổ biến, dữ liệu sản phẩm và đơn hàng được lưu trong database, phân quyền admin giúp tách khu vực quản trị khỏi người dùng thường.

Một số điểm cần tiếp tục hoàn thiện gồm tích hợp thanh toán thật, gửi email xác nhận đơn hàng, bổ sung đánh giá sản phẩm đầy đủ, tối ưu tìm kiếm theo nhiều thuộc tính và triển khai database production như PostgreSQL hoặc MySQL.

# CHƯƠNG VI: KẾT LUẬN

## 6.1 Tổng kết nội dung báo cáo

Báo cáo đã trình bày quá trình phân tích và thiết kế hệ thống website bán phụ kiện máy tính gaming trực tuyến GearZone. Nội dung gồm khảo sát hệ thống tương tự, xác định yêu cầu, thiết kế UML, thiết kế cơ sở dữ liệu, mô tả quá trình thực hiện website và kiểm thử các chức năng chính.

## 6.2 Đánh giá kết quả đạt được

Đề tài đã xây dựng được website có các chức năng cơ bản của một hệ thống thương mại điện tử: xem sản phẩm, tìm kiếm, chi tiết sản phẩm, giỏ hàng, đặt hàng, đăng nhập/đăng ký và quản trị. Hệ thống có cấu trúc mã nguồn rõ ràng, sử dụng công nghệ hiện đại và có khả năng mở rộng.

## 6.3 Hướng phát triển trong tương lai

Trong tương lai, hệ thống có thể phát triển thêm các chức năng:

- Tích hợp cổng thanh toán VNPay, MoMo hoặc VietQR tự động.
- Gửi email xác nhận đơn hàng và thông báo trạng thái vận chuyển.
- Xây dựng hệ thống đánh giá, bình luận và hỏi đáp sản phẩm.
- Thêm chức năng so sánh sản phẩm theo thông số kỹ thuật.
- Tối ưu SEO, tốc độ tải trang và triển khai lên cloud.
- Chuyển database sang PostgreSQL/MySQL cho môi trường production.

## 6.4 Kết luận chung

Đề tài GearZone giúp em vận dụng kiến thức phân tích thiết kế hướng đối tượng, cơ sở dữ liệu và lập trình web vào một bài toán thực tế. Sản phẩm tuy còn có thể tiếp tục hoàn thiện nhưng đã đáp ứng các yêu cầu cơ bản của một website bán phụ kiện gaming trực tuyến và là nền tảng tốt để phát triển thêm trong tương lai.

# TÀI LIỆU THAM KHẢO

1. Tài liệu môn Phân tích thiết kế hướng đối tượng.
2. Next.js Documentation, https://nextjs.org/docs
3. Prisma Documentation, https://www.prisma.io/docs
4. React Documentation, https://react.dev
5. TailwindCSS Documentation, https://tailwindcss.com/docs
6. Tài liệu khảo sát các website GearVN, Phong Vũ, HACOM.
'''


def setup_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    styles = doc.styles
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(13)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(0, 0, 0)
    styles["Heading 3"].font.size = Pt(13)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor(0, 0, 0)


def add_table_from_md(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    # Drop markdown separator row.
    data = [rows[0]] + rows[2:]
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(data):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text.strip(), bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, "D9EAF7")


def add_code_block(doc: Document, code: str, caption: str | None = None) -> None:
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
    for line in code.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8)


def render_markdown_to_docx(doc: Document, md: str) -> None:
    lines = md.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    pending_caption: str | None = None
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                in_code = False
                caption = pending_caption
                add_code_block(doc, "\n".join(code_lines), caption if code_lang == "mermaid" else None)
                pending_caption = None
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("Hình "):
            pending_caption = line.strip()
            # Caption is printed before the following diagram if any; otherwise as normal paragraph.
            if i + 1 >= len(lines) or not lines[i + 1].startswith("```"):
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                pending_caption = None
            i += 1
            continue

        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            add_table_from_md(doc, rows)
            continue

        if not line.strip() or line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_heading(text, level=level)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
            i += 1
            continue

        p = doc.add_paragraph(line.strip())
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
        i += 1


def main() -> None:
    md = markdown_content()
    MD_PATH.write_text(md, encoding="utf-8")

    doc = Document()
    setup_doc(doc)
    add_title_page(doc)
    add_assignment_page(doc)
    add_score_page(doc)
    body_start = md.index("# LỜI CẢM ƠN")
    render_markdown_to_docx(doc, md[body_start:])

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "GVHD: .....................................        SV: ....................................."
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(DOCX_PATH)
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
